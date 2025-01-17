import os
from typing import List, Union
from pinecone import Pinecone, ServerlessSpec
import cohere
import pandas as pd
try:
    from docx import Document
except ImportError:
    print("Warning: python-docx not installed correctly. DOCX support will be limited.")
    Document = None
import PyPDF2

class VectorDatabase:
    def __init__(self):
        self.co = cohere.ClientV2(api_key=os.getenv("COHERE_API_KEY"))
        self.pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        self.index_name = "coffeeshop-index"
        self.model_name = 'embed-english-light-v3.0'
        self.dimension = 384
        
        try:
            try:
                existing_index = self.pc.describe_index(self.index_name)
                if existing_index.dimension != self.dimension:
                    self.pc.delete_index(self.index_name)
            except:
                pass

            self.index = self.pc.create_index(
                name=self.index_name,
                dimension=self.dimension,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
        except Exception:
            self.index = self.pc.Index(self.index_name)

    def read_file(self, file_path: str) -> str:
        ext = file_path.split('.')[-1].lower()
        
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif ext == 'pdf':
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            return text
        
        elif ext == 'docx':
            if Document is None:
                raise ImportError("python-docx is not properly installed")
            doc = Document(file_path)
            return ' '.join([paragraph.text for paragraph in doc.paragraphs])
        
        elif ext == 'csv':
            df = pd.read_csv(file_path)
            return df.to_string()
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def read_file_content(self, file_content: bytes, filename: str) -> str:
        ext = filename.split('.')[-1].lower()
        
        if ext == 'txt':
            return file_content.decode('utf-8')
        
        elif ext == 'pdf':
            text = ""
            pdf_reader = PyPDF2.PdfReader(file_content)
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        
        elif ext == 'docx':
            if Document is None:
                raise ImportError("python-docx is not properly installed")
            from io import BytesIO
            doc = Document(BytesIO(file_content))
            return ' '.join([paragraph.text for paragraph in doc.paragraphs])
        
        elif ext == 'csv':
            from io import StringIO
            df = pd.read_csv(StringIO(file_content.decode('utf-8')))
            return df.to_string()
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def get_embedding(self, text: Union[str, List[str]]):
        texts = [text] if isinstance(text, str) else text
        response = self.co.embed(
            texts=texts,
            model=self.model_name,
            input_type='search_document',
            embedding_types=['float']
        )
        return [list(e) for e in response.embeddings.float_]

    def delete_file(self, file_path: str):
        try:
            os.remove(file_path)
            print(f"File {file_path} deleted successfully.")
        except Exception as e:
            print(f"Error deleting file {file_path}: {e}")

    def add_to_database(self, file_path: str, chunk_size: int = 1000):
        content = self.read_file(file_path)
        chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
        embeddings = self.get_embedding(chunks)
        
        if not embeddings or not all(isinstance(e, list) and all(isinstance(v, float) for v in e) for e in embeddings):
            raise ValueError("Invalid embedding format")
        
        vectors = [{
            'id': f"{file_path}_{i}",
            'values': embedding,
            'metadata': {'text': chunk, 'source': file_path}
        } for i, (chunk, embedding) in enumerate(zip(chunks, embeddings))]
        
        self.index.upsert(vectors=vectors)
        self.delete_file(file_path)

    def add_content_to_database(self, content: str, source_id: str, chunk_size: int = 1000):
        chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
        embeddings = self.get_embedding(chunks)
        
        if not embeddings or not all(isinstance(e, list) and all(isinstance(v, float) for v in e) for e in embeddings):
            raise ValueError("Invalid embedding format")
        
        vectors = [{
            'id': f"{source_id}_{i}",
            'values': embedding,
            'metadata': {'text': chunk, 'source': source_id}
        } for i, (chunk, embedding) in enumerate(zip(chunks, embeddings))]
        
        self.index.upsert(vectors=vectors)

    def query_database(self, query: str, top_k: int = 3):
        query_embedding = self.get_embedding(query)[0]
        results = self.index.query(
            vector=query_embedding,
            top_k=top_k,
            include_values=False,
            include_metadata=True
        )
        
        return [{
            'text': match.metadata['text'],
            'source': match.metadata.get('source', 'Unknown'),
            'score': match.score
        } for match in results.matches]
